"""Montagem do .docx — camadas 1 + 2 + 3 no layout do laudo real.

Nenhum texto nasce aqui. O documento é a soma de:
  - camada 1, como o perito ditou e revisou;
  - camada 2, transcrita do laudo SB 1252/2019;
  - camada 3, derivada por regra e **confirmada** na tela anterior.

Onde falta redação transcrita, o documento carrega o marcador
``[PENDENTE: ...]`` à vista, em vermelho. Uma minuta com lacuna assinalada é
honesta; uma minuta com lacuna preenchida por semelhança, não.
"""

from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config.schema import Exame, Secao
from core import derivados as camada3
from core import numeros
from core import quesitos as camada1_quesitos
from core import templates as texto_fixo

FONTE = "Times New Roman"
CORPO = Pt(12)
VERMELHO = RGBColor(0xC0, 0x00, 0x00)
MARCADOR = "[PENDENTE:"


def _paginas(derivados: dict) -> str:
    """Contagem por extenso quando o perito informou; vazio quando não."""
    return numeros.paginas_por_extenso(derivados.get(camada3.CHAVE_PAGINAS, ""))


def _insere_contagem_automatica(paragrafo) -> None:
    """Campo NUMPAGES: o editor de texto conta as páginas ao abrir o arquivo.

    Paginação não existe nesta montagem — quem pagina é o Word. Em vez de pedir
    ao perito que conte e volte, o número entra como campo e sai certo sozinho.
    """
    corrida = paragrafo.add_run()
    abertura = OxmlElement("w:fldChar")
    abertura.set(qn("w:fldCharType"), "begin")
    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = " NUMPAGES "
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    provisorio = OxmlElement("w:t")
    provisorio.text = "2"
    fechamento = OxmlElement("w:fldChar")
    fechamento.set(qn("w:fldCharType"), "end")
    for elemento in (abertura, instrucao, separador, provisorio, fechamento):
        corrida._r.append(elemento)


def _paragrafo_fecho(documento: Document, ctx) -> None:
    """Fecho do laudo, com a contagem de páginas por extenso ou automática."""
    modelo = ctx.texto("FECHO", "")
    if not modelo:
        return
    extenso = _paginas(ctx.derivados)
    if extenso:
        _paragrafo(documento, _formata(modelo, {"paginas_extenso": extenso}))
        return

    antes, _, depois = modelo.partition("{paginas_extenso}")
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo.paragraph_format.space_after = Pt(6)
    paragrafo.add_run(antes)
    _insere_contagem_automatica(paragrafo)
    paragrafo.add_run(depois)


def _configura(documento: Document) -> None:
    estilo = documento.styles["Normal"]
    estilo.font.name = FONTE
    estilo.font.size = CORPO
    for secao in documento.sections:
        secao.top_margin = Cm(2.5)
        secao.bottom_margin = Cm(2.5)
        secao.left_margin = Cm(3.0)
        secao.right_margin = Cm(2.5)


def _paragrafo(
    documento: Document,
    texto: str,
    *,
    alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
    negrito: bool = False,
    espaco_antes: int = 0,
):
    """Parágrafo que destaca em vermelho os trechos ainda pendentes."""
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = alinhamento
    paragrafo.paragraph_format.space_before = Pt(espaco_antes)
    paragrafo.paragraph_format.space_after = Pt(6)

    restante = texto
    while MARCADOR in restante:
        antes, _, depois = restante.partition(MARCADOR)
        if antes:
            corrida = paragrafo.add_run(antes)
            corrida.bold = negrito
        pendencia, _, restante = depois.partition("]")
        alerta = paragrafo.add_run(f"{MARCADOR}{pendencia}]")
        alerta.bold = True
        alerta.font.color.rgb = VERMELHO
    if restante:
        corrida = paragrafo.add_run(restante)
        corrida.bold = negrito
    return paragrafo


def _titulo_secao(documento: Document, texto: str) -> None:
    _paragrafo(
        documento,
        texto,
        alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
        negrito=True,
        espaco_antes=12,
    )


@dataclass
class Contexto:
    """Tudo que uma seção pode precisar, sem assinaturas gigantes."""

    exame: Exame
    admin: dict
    colecoes: dict
    derivados: dict
    imagens: list[dict]
    quesitos: list[str]
    respostas_quesitos: dict[str, str]

    @property
    def boiler(self):
        return texto_fixo.boilerplate(self.exame)

    def texto(self, nome: str, padrao: str = "") -> str:
        return texto_fixo.texto(self.exame, nome, padrao)

    def peritos(self) -> list[dict]:
        """Signatários. Um grupo repetível, ou o campo único, nessa ordem."""
        for grupo in self.exame.grupos_admin:
            registrados = self.admin.get(grupo.chave)
            if isinstance(registrados, list) and registrados:
                return [p for p in registrados if any(str(v).strip() for v in p.values())]
        unico = {
            "perito_designado": self.admin.get("perito_designado", ""),
            "matricula": self.admin.get("matricula", ""),
            "classe_perito": self.admin.get("classe_perito", ""),
        }
        return [unico] if unico["perito_designado"] else []


def _cabecalho(documento: Document, ctx: Contexto, secao: Secao) -> None:
    for linha in ctx.texto("CABECALHO", ()):
        _paragrafo(documento, linha, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True)

    documento.add_paragraph()
    for rotulo, chave in (("DEMANDA", "numero_demanda"), ("LAUDO N°", "numero_laudo")):
        valor = ctx.admin.get(chave, "")
        if valor:
            _paragrafo(
                documento,
                f"{rotulo}        {valor}",
                alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
                negrito=True,
            )

    documento.add_paragraph()
    for constante in ("TITULO", "SUBTITULO"):
        linha = ctx.texto(constante)
        if linha:
            _paragrafo(
                documento, linha, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True
            )
    documento.add_paragraph()


def _campos_do_preambulo(ctx: Contexto) -> dict:
    """Valores disponíveis ao texto do preâmbulo, por nome de campo."""
    valores = {
        chave: valor for chave, valor in ctx.admin.items() if not isinstance(valor, list)
    }
    valores["data_exame_extenso"] = numeros.data_por_extenso(ctx.admin.get("data_exame", ""))
    valores["data_documento"] = numeros.data_curta(ctx.admin.get("data_documento", ""))
    valores["data_documento_extenso"] = numeros.data_por_extenso(
        ctx.admin.get("data_documento", "")
    )
    valores["data_recebimento_extenso"] = numeros.data_por_extenso(
        ctx.admin.get("data_recebimento") or ctx.admin.get("data_exame", "")
    )
    valores["data_realizacao"] = numeros.data_curta(ctx.admin.get("data_realizacao", ""))

    peritos = ctx.peritos()
    nomes = [str(p.get("perito_designado", "")).strip() for p in peritos]
    nomes = [n for n in nomes if n]
    if len(nomes) > 1:
        valores["peritos_designados"] = ", ".join(nomes[:-1]) + " e " + nomes[-1]
        valores["perito_ou_peritos"] = "os PERITOS CRIMINAIS"
    else:
        valores["peritos_designados"] = nomes[0] if nomes else ""
        valores["perito_ou_peritos"] = "o(a) PERITO(A) CRIMINAL"
    valores.setdefault("perito_designado", valores["peritos_designados"])
    return valores


def _preambulo(documento: Document, ctx: Contexto, secao: Secao) -> None:
    modelo = ctx.texto("PREAMBULO")
    if modelo:
        _paragrafo(documento, _formata(modelo, _campos_do_preambulo(ctx)))


def _texto_fixo(documento: Document, ctx: Contexto, secao: Secao) -> None:
    """Seção de texto do template, parametrizada pelos campos do caso."""
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)
    valores = _campos_do_preambulo(ctx)
    for nome in [c.strip() for c in secao.chave.split(",") if c.strip()]:
        modelo = ctx.texto(nome)
        if modelo:
            _paragrafo(documento, _formata(modelo, valores))


def _objetos(documento: Document, ctx: Contexto, secao: Secao) -> None:
    """Descrição dos itens examinados, um parágrafo por item."""
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)

    chave = secao.chave or "materiais"
    itens = ctx.colecoes.get(chave, [])
    varios = len(itens) > 1
    for indice, item in enumerate(itens, start=1):
        descricao = ctx.derivados.get(
            f"{camada3.PREFIXO_MATERIAL}{indice}"
        ) or camada3.descricao_material(item)
        prefixo = f"{chr(ord('a') + indice - 1)}) " if varios else ""
        _paragrafo(documento, f"{prefixo}{descricao}")

    if not ctx.exame.imagens_em_apendice:
        _imagens_no_corpo(documento, ctx, chave)


def _imagens_no_corpo(documento: Document, ctx: Contexto, colecao: str) -> None:
    if not ctx.imagens:
        return
    documento.add_paragraph()
    for imagem in ctx.imagens:
        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo.add_run().add_picture(io.BytesIO(imagem["dados"]), width=Cm(10))
        legenda = imagem.get("legenda") or ctx.texto("LEGENDA_FOTO", "")
        if legenda:
            _paragrafo(documento, legenda, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)


def _exames(documento: Document, ctx: Contexto, secao: Secao) -> None:
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)
    narrativa = ctx.derivados.get(camada3.CHAVE_EXAMES) or camada3.texto_exames(ctx.colecoes)
    if narrativa:
        _paragrafo(documento, narrativa)


def _resultados(documento: Document, ctx: Contexto, secao: Secao) -> None:
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)
    numero = secao.titulo.split(".")[0].strip() if secao.titulo else ""
    for ordem, item in enumerate(
        camada3.resultados_obtidos(ctx.colecoes, ctx.derivados), start=1
    ):
        rotulo = f"{numero}.{ordem}. {item['titulo']}" if numero else item["titulo"]
        _paragrafo(
            documento,
            rotulo,
            alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
            negrito=True,
            espaco_antes=6,
        )
        _paragrafo(documento, item["texto"])


def _conclusao(documento: Document, ctx: Contexto, secao: Secao) -> None:
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)
    resultados = ctx.derivados.get(camada3.CHAVE_CONCLUSAO) or camada3.conclusao(ctx.colecoes)[0]
    modelo = ctx.texto("CONCLUSAO", "{resultados}")
    _paragrafo(documento, _formata(modelo, {"resultados": resultados.rstrip(".")}))


def _quesitos(documento: Document, ctx: Contexto, secao: Secao) -> None:
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)
    abertura = ctx.texto("ABERTURA_QUESITOS")
    if abertura:
        documento.add_paragraph()
        _paragrafo(documento, abertura)

    perguntas = ctx.quesitos or list(ctx.texto("QUESITOS_DA_REQUISICAO_MODELO", ()))
    for quesito in camada1_quesitos.montar(
        perguntas, ctx.colecoes, ctx.derivados, ctx.respostas_quesitos
    ):
        _paragrafo(documento, f"{quesito.numero} – {quesito.pergunta}", espaco_antes=6)
        _paragrafo(documento, "R – " + quesito.resposta)


def _referencias(documento: Document, ctx: Contexto, secao: Secao) -> None:
    if secao.titulo:
        _titulo_secao(documento, secao.titulo)
    for referencia in camada3.referencias(ctx.colecoes):
        _paragrafo(documento, referencia)


def _fecho(documento: Document, ctx: Contexto, secao: Secao) -> None:
    documento.add_paragraph()
    _paragrafo_fecho(documento, ctx)


def _assinatura(documento: Document, ctx: Contexto, secao: Secao) -> None:
    documento.add_paragraph()
    assinatura = ctx.texto("ASSINATURA", "")
    cargo = ctx.texto("CARGO", "")
    for perito in ctx.peritos():
        if assinatura:
            _paragrafo(documento, assinatura, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
        nome = str(perito.get("perito_designado", "")).strip().upper()
        if nome:
            _paragrafo(
                documento, nome, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, negrito=True
            )
        if cargo:
            _paragrafo(documento, cargo, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
        classe = str(perito.get("classe_perito", "")).strip()
        matricula = str(perito.get("matricula", "")).strip()
        rodape = f"{classe} – Matrícula: {matricula}" if classe else f"Matrícula: {matricula}"
        _paragrafo(documento, rodape, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
        documento.add_paragraph()


def _apendice(documento: Document, ctx: Contexto, secao: Secao) -> None:
    """Apêndice fotográfico: página própria, imagens numeradas com legenda."""
    if not ctx.imagens:
        return
    documento.add_page_break()
    _paragrafo(
        documento,
        secao.titulo or "APÊNDICE FOTOGRÁFICO",
        alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
        negrito=True,
    )
    documento.add_paragraph()
    for numero, imagem in enumerate(ctx.imagens, start=1):
        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo.add_run().add_picture(io.BytesIO(imagem["dados"]), width=Cm(8))
        legenda = imagem.get("legenda") or f"Imagem {numero:02d}"
        _paragrafo(documento, legenda, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)


#: Ordem clássica do laudo de substância, usada quando o exame não declara a sua.
SECOES_PADRAO = (
    Secao("cabecalho"),
    Secao("preambulo"),
    Secao("texto", "1. HISTÓRICO", "HISTORICO,HISTORICO_FECHO"),
    Secao("objetos", "2. IDENTIFICAÇÃO E DESCRIÇÃO DO MATERIAL", "materiais"),
    Secao("exames", "3. EXAMES REALIZADOS"),
    Secao("resultados", "4. RESULTADOS OBTIDOS"),
    Secao("conclusao", "5. CONCLUSÃO"),
    Secao("quesitos"),
    Secao("referencias", "6. REFERÊNCIAS"),
    Secao("fecho"),
    Secao("assinatura"),
)

MONTADORES = {
    "cabecalho": _cabecalho,
    "preambulo": _preambulo,
    "texto": _texto_fixo,
    "objetos": _objetos,
    "exames": _exames,
    "resultados": _resultados,
    "conclusao": _conclusao,
    "quesitos": _quesitos,
    "referencias": _referencias,
    "fecho": _fecho,
    "assinatura": _assinatura,
    "apendice": _apendice,
}


def _formata(modelo: str, valores: dict) -> str:
    """Preenche o template tolerando marcador que o caso não tem."""

    class _Faltante(dict):
        def __missing__(self, chave):  # noqa: D105
            return ""

    try:
        return modelo.format_map(_Faltante(valores))
    except (IndexError, ValueError):
        return modelo


def montar(
    admin: dict,
    colecoes: dict[str, list[dict]],
    derivados: dict,
    imagens: list[dict] | None = None,
    quesitos: list[str] | None = None,
    respostas_quesitos: dict[str, str] | None = None,
    exame: Exame | None = None,
) -> Document:
    """Monta o .docx percorrendo as seções que o TIPO DE EXAME declara."""
    if exame is None:
        from config.exams import obter_exame

        exame = obter_exame(texto_fixo.PADRAO)
    ctx = Contexto(
        exame=exame,
        admin=admin,
        colecoes=colecoes,
        derivados=derivados,
        imagens=imagens or [],
        quesitos=quesitos or [],
        respostas_quesitos=respostas_quesitos or {},
    )
    documento = Document()
    _configura(documento)
    for secao in (exame.secoes if exame and exame.secoes else SECOES_PADRAO):
        montador = MONTADORES.get(secao.tipo)
        if montador is not None:
            montador(documento, ctx, secao)
    return documento


def em_bytes(documento: Document) -> bytes:
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def pendencias_do_texto(
    admin: dict,
    colecoes: dict[str, list[dict]],
    derivados: dict,
    quesitos: list[str] | None = None,
    respostas_quesitos: dict[str, str] | None = None,
) -> list[str]:
    """Marcadores [PENDENTE: ...] que apareceriam no documento."""
    perguntas = quesitos or []
    pedacos = [
        *camada3.referencias(colecoes),
        *(
            q.resposta
            for q in camada1_quesitos.montar(
                perguntas, colecoes, derivados, respostas_quesitos or {}
            )
        ),
        derivados.get(camada3.CHAVE_NATUREZA) or camada3.natureza(colecoes),
        derivados.get(camada3.CHAVE_PROSCRICAO) or camada3.proscricao(colecoes),
        *(s["texto"] for s in camada3.resultados_obtidos(colecoes, derivados)),
        *(
            derivados.get(f"{camada3.PREFIXO_MATERIAL}{i}")
            or camada3.descricao_material(m)
            for i, m in enumerate(colecoes.get("materiais", []), start=1)
        ),
    ]
    encontradas: list[str] = []
    for pedaco in pedacos:
        restante = str(pedaco)
        while MARCADOR in restante:
            _, _, depois = restante.partition(MARCADOR)
            pendencia, _, restante = depois.partition("]")
            texto = pendencia.strip()
            if texto not in encontradas:
                encontradas.append(texto)
    return encontradas
